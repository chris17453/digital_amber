#!/usr/bin/env python3
"""Build Amazon KDP format from markdown files."""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def read_file(path):
    """Read file content with UTF-8 encoding."""
    return Path(path).read_text(encoding='utf-8')

def create_kindle_docx():
    """Create a DOCX file formatted for Amazon KDP."""
    doc = Document()
    
    # Set up document margins for KDP (6" x 9" format)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.625)
        section.right_margin = Inches(0.625)
    
    # Create custom styles
    styles = doc.styles
    
    # Title page style
    title_style = styles.add_style('BookTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Times New Roman'
    title_font.size = Pt(24)
    title_font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(24)
    
    # Subtitle style
    subtitle_style = styles.add_style('BookSubtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Times New Roman'
    subtitle_font.size = Pt(16)
    subtitle_font.italic = True
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_style.paragraph_format.space_after = Pt(12)
    
    # Author style
    author_style = styles.add_style('BookAuthor', WD_STYLE_TYPE.PARAGRAPH)
    author_font = author_style.font
    author_font.name = 'Times New Roman'
    author_font.size = Pt(14)
    author_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_style.paragraph_format.space_after = Pt(24)
    
    # Chapter heading style
    chapter_style = styles.add_style('ChapterHeading', WD_STYLE_TYPE.PARAGRAPH)
    chapter_font = chapter_style.font
    chapter_font.name = 'Times New Roman'
    chapter_font.size = Pt(18)
    chapter_font.bold = True
    chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chapter_style.paragraph_format.space_before = Pt(36)
    chapter_style.paragraph_format.space_after = Pt(24)
    
    # Body text style
    body_style = styles.add_style('BookBody', WD_STYLE_TYPE.PARAGRAPH)
    body_font = body_style.font
    body_font.name = 'Times New Roman'
    body_font.size = Pt(11)
    body_style.paragraph_format.line_spacing = 1.15
    body_style.paragraph_format.space_after = Pt(6)
    body_style.paragraph_format.first_line_indent = Inches(0.25)
    
    return doc

def add_title_page(doc):
    """Add title page to document."""
    # Title
    title_p = doc.add_paragraph('Digital Amber', style='BookTitle')
    
    # Subtitle
    subtitle_p = doc.add_paragraph('When Consciousness Becomes Code', style='BookSubtitle')
    
    # Add some space
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # Tagline
    tagline_p = doc.add_paragraph('A Speculative Exploration of AI, Identity, and the Future of Mind', style='BookSubtitle')
    
    # Add more space
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    # Author
    author_p = doc.add_paragraph('By Charles Watkins', style='BookAuthor')
    
    # Page break
    doc.add_page_break()

def process_markdown_to_docx(content, doc):
    """Convert markdown content to Word document paragraphs."""
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph('')
            continue
            
        # Handle headers
        if line.startswith('# '):
            doc.add_paragraph(line[2:], style='ChapterHeading')
        elif line.startswith('## '):
            p = doc.add_paragraph(line[3:])
            p.style.font.size = Pt(14)
            p.style.font.bold = True
            p.style.paragraph_format.space_before = Pt(18)
            p.style.paragraph_format.space_after = Pt(12)
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:])
            p.style.font.size = Pt(12)
            p.style.font.bold = True
            p.style.paragraph_format.space_before = Pt(12)
            p.style.paragraph_format.space_after = Pt(6)
        else:
            # Regular paragraph
            # Process bold and italic
            line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)  # Remove bold markdown (Word will handle formatting)
            line = re.sub(r'\*(.+?)\*', r'\1', line)      # Remove italic markdown
            
            p = doc.add_paragraph(line, style='BookBody')

def build_kindle_book():
    """Build the complete Kindle book."""
    output_dir = Path("dist")
    story_dir = Path("story")
    
    output_dir.mkdir(exist_ok=True)
    
    # Create document
    doc = create_kindle_docx()
    
    # Add title page
    add_title_page(doc)
    
    # Add copyright page
    doc.add_paragraph('Copyright © 2025 Charles Watkins', style='BookBody')
    doc.add_paragraph('All rights reserved', style='BookBody')
    doc.add_paragraph('First Digital Edition', style='BookBody')
    doc.add_page_break()
    
    # Add dedication if exists
    if (story_dir / "dedication.md").exists():
        content = read_file(story_dir / "dedication.md")
        process_markdown_to_docx(content, doc)
        doc.add_page_break()
    
    # Add foreword if exists
    if (story_dir / "foreword.md").exists():
        content = read_file(story_dir / "foreword.md")
        process_markdown_to_docx(content, doc)
        doc.add_page_break()
    
    # Add all chapters in order
    chapters = [
        "chapter_1.md", "chapter_2.md", "chapter_3.md", "chapter_4.md", "chapter_5.md",
        "chapter_6.md", "chapter_7.md", "chapter_8.md", "chapter_9.md", "chapter_10.md",
        "chapter_11.md", "chapter_12.md", "chapter_13.md", "chapter_14.md", "chapter_15.md",
        "chapter_16.md", "chapter_17.md", "chapter_18.md", "chapter_19.md", "chapter_20.md",
        "chapter_21.md", "chapter_22.md", "chapter_23.md", "chapter_24.md"
    ]
    
    for chapter_file in chapters:
        chapter_path = story_dir / chapter_file
        if chapter_path.exists():
            content = read_file(chapter_path)
            process_markdown_to_docx(content, doc)
            doc.add_page_break()
    
    # Add epilogue if exists
    if (story_dir / "epilogue.md").exists():
        content = read_file(story_dir / "epilogue.md")
        process_markdown_to_docx(content, doc)
        doc.add_page_break()
    
    # Add acknowledgments if exists
    if (story_dir / "acknowledgements.md").exists():
        content = read_file(story_dir / "acknowledgements.md")
        process_markdown_to_docx(content, doc)
        doc.add_page_break()
    
    # Add about the author if exists
    if (story_dir / "about_the_author.md").exists():
        content = read_file(story_dir / "about_the_author.md")
        process_markdown_to_docx(content, doc)
    
    # Save the document
    output_file = output_dir / "digital_amber_kindle.docx"
    doc.save(output_file)
    
    print(f"Kindle format created: {output_file}")
    print("Ready for upload to Amazon KDP!")

if __name__ == "__main__":
    build_kindle_book()